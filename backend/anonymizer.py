from fastapi import FastAPI, File, UploadFile, Request, HTTPException, Depends, status, Form
from fastapi.responses import JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
from io import BytesIO
import pandas as pd
import docx
import re
import logging
import uvicorn
import traceback
from datetime import datetime, timedelta
from typing import Optional
from passlib.context import CryptContext
from jose import JWTError, jwt
from fastapi.security import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from pydantic import BaseModel
from database import engine, get_db, SessionLocal
import models
from sqlalchemy.orm import Session

# Password hashing
pwd_context = CryptContext(schemes=["bcrypt"], deprecated="auto")

# JWT settings
SECRET_KEY = "your-secret-key-keep-it-secret"  # Change this in production
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30

# Create database tables
models.Base.metadata.create_all(bind=engine)

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

app = FastAPI()

# Configure CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# OAuth2 scheme
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# Pydantic models
class Token(BaseModel):
    access_token: str
    token_type: str

class TokenData(BaseModel):
    username: Optional[str] = None

class UserBase(BaseModel):
    username: str
    email: str

class UserCreate(UserBase):
    password: str

class User(UserBase):
    id: int
    created_at: datetime
    class Config:
        orm_mode = True

# Helper functions
def verify_password(plain_password, hashed_password):
    return pwd_context.verify(plain_password, hashed_password)

def get_password_hash(password):
    return pwd_context.hash(password)

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire})
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt

def get_user(db: Session, username: str):
    return db.query(models.User).filter(models.User.username == username).first()

def get_user_by_email(db: Session, email: str):
    return db.query(models.User).filter(models.User.email == email).first()

async def get_current_user(token: str = Depends(oauth2_scheme), db: Session = Depends(get_db)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        token_data = TokenData(username=username)
    except JWTError:
        raise credentials_exception
    user = get_user(db, username=token_data.username)
    if user is None:
        raise credentials_exception
    return user

# Authentication endpoints
@app.post("/register", response_model=Token)
async def register(user: UserCreate, db: Session = Depends(get_db)):
    db_user = get_user(db, username=user.username)
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    db_user = get_user_by_email(db, email=user.email)
    if db_user:
        raise HTTPException(status_code=400, detail="Email already registered")
    
    hashed_password = get_password_hash(user.password)
    db_user = models.User(
        username=user.username,
        email=user.email,
        password=hashed_password
    )
    db.add(db_user)
    db.commit()
    db.refresh(db_user)
    
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/login", response_model=Token)
async def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = get_user(db, form_data.username)
    if not user or not verify_password(form_data.password, user.password):
        raise HTTPException(
            status_code=401,
            detail="Incorrect username or password",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {"access_token": access_token, "token_type": "bearer"}

@app.get("/me", response_model=User)
async def read_users_me(current_user: User = Depends(get_current_user)):
    return current_user

# ---------------- Helper Functions ----------------

def extract_text_from_file(uploaded_file: UploadFile):
    try:
        content_type = uploaded_file.content_type
        logger.info(f"Processing file with content type: {content_type}")
        
        # Read file content
        file_bytes = uploaded_file.file.read()
        if not file_bytes:
            return "Error: Empty file", None
            
        file_stream = BytesIO(file_bytes)
        
        # Handle text files
        if content_type == "text/plain" or uploaded_file.filename.lower().endswith('.txt'):
            try:
                text = file_bytes.decode("utf-8")
                if not text.strip():
                    return "Error: File is empty", None
                return text, "txt"
            except UnicodeDecodeError:
                try:
                    text = file_bytes.decode("latin-1")
                    if not text.strip():
                        return "Error: File is empty", None
                    return text, "txt"
                except Exception as e:
                    logger.error(f"Error decoding text file: {str(e)}")
                    return f"Error: Unable to decode text file - {str(e)}", None
        
        # Handle Excel files
        elif (content_type in [
            "application/vnd.ms-excel",
            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            "application/octet-stream"
        ] or uploaded_file.filename.lower().endswith(('.xlsx', '.xls'))):
            try:
                # Try openpyxl first (for newer Excel files)
                try:
                    df = pd.read_excel(file_stream, engine="openpyxl")
                except Exception:
                    # If openpyxl fails, try xlrd (for older Excel files)
                    file_stream.seek(0)
                    df = pd.read_excel(file_stream, engine="xlrd")
                
                if df.empty:
                    return "Error: Excel file is empty", None
                    
                # Convert all data to strings and handle NaN values
                for col in df.columns:
                    df[col] = df[col].fillna("").astype(str)
                return df, "excel"
                
            except Exception as e:
                logger.error(f"Error processing Excel file: {str(e)}")
                return f"Error: Unable to process Excel file - {str(e)}", None
        
        # Handle Word documents
        elif (content_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" 
              or uploaded_file.filename.lower().endswith('.docx')):
            try:
                doc = docx.Document(file_stream)
                
                # Extract text from paragraphs and tables
                text_parts = []
                
                # Get text from paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                
                # Get text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                if not text_parts:
                    return "Error: No readable text found in the document", None
                    
                return "\n".join(text_parts), "docx"
                
            except Exception as e:
                logger.error(f"Error processing Word document: {str(e)}")
                return f"Error: Unable to process Word document - {str(e)}", None
        
        else:
            return f"Error: Unsupported file type - {content_type}", None
            
    except Exception as e:
        logger.error(f"Error in extract_text_from_file: {str(e)}")
        traceback.print_exc()
        return f"Error: File processing failed - {str(e)}", None

def anonymize_text(text):
    try:
        if isinstance(text, pd.DataFrame):
            # Handle DataFrame (Excel files)
            # Define sensitive columns that should be dropped
            sensitive_columns = [
                'Patient ID', 'Patient ID 1', 'Attendance Number', 'Patient Name', 
                'NHS Number', 'Hospital Number', 'DOB', 'Date of Birth', 'Age',
                'Phone', 'Mobile', 'Email', 'Address', 'Postcode', 'Name'
            ]
            
            # Case-insensitive column matching
            columns_to_drop = []
            for col in text.columns:
                if any(sensitive.lower() in col.lower() for sensitive in sensitive_columns):
                    columns_to_drop.append(col)
            
            # Drop sensitive columns if they exist
            if columns_to_drop:
                text = text.drop(columns=columns_to_drop)
            
            # Apply anonymization to all remaining text cells
            for col in text.columns:
                text[col] = text[col].apply(lambda x: anonymize_text(str(x)) if pd.notna(x) else '')
            
            return text

        if isinstance(text, str):
            # Handle text content
            patterns = {
                # Patient Information
                r'\b(?:Mr\.?|Mrs\.?|Ms\.?|Miss|Dr\.?)\s*[A-Za-z\'-]+(?:\s+[A-Za-z\'-]+){0,2}\b': "[PERSON NAME]",
                r'\b(?:Patient|Name):\s*[A-Za-z\'-]+(?:\s+[A-Za-z\'-]+){0,2}\b': "[PATIENT NAME]",
                r'\b(?:DOB|Date\s+of\s+Birth):\s*\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b': "[DATE OF BIRTH]",
                r'\bAge:\s*\d+\b': "[AGE]",
                r'\bNHS\s*(?:No|Number|#)?\s*:?\s*\d{3}[-\s]*\d{3}[-\s]*\d{4}\b': "[NHS NUMBER]",
                r'\bHospital\s*(?:No|Number|#)?\s*:?\s*\d+\b': "[HOSPITAL NUMBER]",
                r'\bPatient\s*ID\s*:?\s*\d+\b': "[PATIENT ID]",
                
                # Contact Information
                r'\b(?:Phone|Tel|Mobile|Contact)(?:\s*\d{2,}){2,}\b': "[PHONE NUMBER]",
                r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b': "[EMAIL]",
                r'\b(?:Post\s*code|Postcode):\s*[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2}\b': "[POSTCODE]",
                
                # Medical Staff
                r'\b(?:Dr|Doctor|Prof|Professor)\.\s*[A-Za-z\'-]+(?:\s+[A-Za-z\'-]+){0,2}\b': "[DOCTOR NAME]",
                r'\bReported\s*by\s*(?:Dr\.)?\s*[A-Za-z\'-]+(?:\s+[A-Za-z\'-]+){0,2}\b': "[REPORTED BY]",
                r'\bConsultant\s*Radiologist(?:\(s\))?\b': "[CONSULTANT]",
                r'\bREPORT\s*AUTHOR\s*:\s*(?:Dr\.)?\s*[A-Za-z\'-]+(?:\s+[A-Za-z\'-]+){0,2}\b': "[REPORT AUTHOR]",
                
                # Location Information
                r'\b(?:The\s*)?Jack\s*Brignall\s*PET/CT\s*Centre\b': "[MEDICAL CENTER]",
                r'\bCastle\s*Hill\s*Hospital\b': "[HOSPITAL]",
                r'\bCastle\s*Road\b': "[STREET]",
                r'\bCottingham\b': "[CITY]",
                r'\bEast\s*Riding\s*Of\s*Yorkshire\b': "[REGION]",
                r'\b[A-Z]{1,2}\d{1,2}[A-Z]?\s*\d[A-Z]{2}\b': "[POSTAL CODE]",
                
                # Identifiers and Numbers
                r'\b\d{6,}\b': "[ID NUMBER]",
                r'\bFDG\s*Batch\s*Number:\s*\d{2}-\d{4}-\d{2}\b': "[BATCH NUMBER]",
                
                # Signatures and Staff Actions
                r'\bElectronically\s*signed\s*by:\s*(?:Dr\.)?\s*[A-Za-z\'-]+(?:\s+[A-Za-z\'-]+){0,2}\s*\([^)]+\)': "[ELECTRONIC SIGNATURE]",
                r'\bScanned\s*by:\s*[A-Za-z\'-]+,\s*[A-Za-z\'-]+\b': "[SCANNED BY]",
                r'\bFDG\s*Injected\s*by:\s*[A-Za-z\'-]+,\s*[A-Za-z\'-]+\b': "[INJECTED BY]",
                
                # Dates and Times
                r'\b\d{1,2}[-/]\d{1,2}[-/]\d{2,4}\b': "[DATE]",
                r'\b\d{1,2}:\d{2}(?::\d{2})?\s*(?:AM|PM)?\b': "[TIME]",
            }
            
            # Sort patterns by length (longest first) to handle overlapping matches
            sorted_patterns = sorted(patterns.items(), key=lambda x: len(x[0]), reverse=True)
            
            # Apply patterns
            for pattern, replacement in sorted_patterns:
                text = re.sub(pattern, replacement, text, flags=re.IGNORECASE)
            
            return text
            
        return text
        
    except Exception as e:
        logger.error(f"Error in anonymize_text: {str(e)}")
        traceback.print_exc()
        return f"Error: Anonymization failed - {str(e)}"

def create_downloadable_file(data, file_format):
    try:
        output = BytesIO()
        
        if file_format == "txt":
            output.write(data.encode("utf-8"))
            mime_type = "text/plain"
            file_extension = ".txt"
        
        elif file_format == "excel":
            # Convert string representation of list of dicts back to DataFrame
            if isinstance(data, str):
                try:
                    data_list = json.loads(data)
                    df = pd.DataFrame(data_list)
                except Exception as e:
                    logger.error(f"Error converting JSON to DataFrame: {str(e)}")
                    # If conversion fails, create a simple DataFrame with the text
                    df = pd.DataFrame({"Anonymized Content": [data]})
            else:
                df = pd.DataFrame(data)
                
            with pd.ExcelWriter(output, engine="openpyxl") as writer:
                df.to_excel(writer, index=False)
            
            mime_type = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            file_extension = ".xlsx"
        
        elif file_format == "docx":
            doc = docx.Document()
            doc.add_paragraph(data)
            doc.save(output)
            mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            file_extension = ".docx"
        
        output.seek(0)
        return output, mime_type, file_extension
    except Exception as e:
        logger.error(f"Error in create_downloadable_file: {str(e)}")
        traceback.print_exc()
        raise

# ---------------- API Endpoints ----------------

@app.post("/anonymize")
async def anonymize_endpoint(request: Request, file: UploadFile = File(...)):
    try:
        logger.info(f"Processing file: {file.filename}")
        
        # Validate file size (e.g., limit to 10MB)
        file_size = 0
        chunk_size = 1024 * 1024  # 1MB chunks
        content = BytesIO()
        
        while True:
            chunk = await file.read(chunk_size)
            if not chunk:
                break
            file_size += len(chunk)
            if file_size > 10 * 1024 * 1024:  # 10MB limit
                raise HTTPException(status_code=413, detail="File too large. Maximum size is 10MB")
            content.write(chunk)
        
        content.seek(0)
        file.file = content
        
        # Extract text from file
        text, file_format = extract_text_from_file(file)
        
        # Check for extraction errors
        if isinstance(text, str) and text.startswith("Error:"):
            raise HTTPException(status_code=400, detail=text)
        
        if text is None or file_format is None:
            raise HTTPException(status_code=400, detail="Unable to process the uploaded file")
        
        # Anonymize the text
        anonymized_result = anonymize_text(text)
        
        # Check for anonymization errors
        if isinstance(anonymized_result, str) and anonymized_result.startswith("Error:"):
            raise HTTPException(status_code=500, detail=anonymized_result)
        
        # Create downloadable file
        if file_format == "excel":
            # Convert DataFrame to list of dicts for JSON response
            result = {
                "content": anonymized_result.to_dict(orient="records"),
                "format": file_format
            }
        else:
            result = {
                "content": anonymized_result,
                "format": file_format
            }
        
        logger.info(f"Successfully anonymized file: {file.filename}")
        return JSONResponse(content=result)
        
    except HTTPException as he:
        logger.error(f"HTTP error in anonymize endpoint: {str(he.detail)}")
        raise he
    except Exception as e:
        logger.error(f"Error in anonymize endpoint: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An error occurred while processing the file: {str(e)}")

@app.post("/download_anonymized")
async def download_anonymized(request: Request, file_content: str = Form(...), file_format: str = Form(...)):
    try:
        logger.info("Download request received")
        
        if not file_content or not file_format:
            raise HTTPException(status_code=400, detail="Missing required parameters")
        
        # Create downloadable file
        output, mime_type, file_extension = create_downloadable_file(file_content, file_format)
        
        if isinstance(output, str) and output.startswith("Error:"):
            raise HTTPException(status_code=500, detail=output)
        
        # Generate a timestamp for the filename
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"anonymized_{timestamp}{file_extension}"
        
        logger.info(f"Sending anonymized file: {filename}")
        return StreamingResponse(
            output,
            media_type=mime_type,
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
        
    except HTTPException as he:
        logger.error(f"HTTP error in download endpoint: {str(he.detail)}")
        raise he
    except Exception as e:
        logger.error(f"Error in download endpoint: {str(e)}")
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An error occurred while creating the download: {str(e)}")

# Test endpoint to verify the server is running
@app.get("/test")
def test():
    return {"status": "ok", "message": "Anonymizer API is running"}

# Root endpoint
@app.get("/")
def root():
    return {"message": "MediXscan Anonymizer API"}

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8005)
